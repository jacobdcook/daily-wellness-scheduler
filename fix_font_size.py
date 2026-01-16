#!/usr/bin/env python3
"""
Fix font size on project name to match other project titles
"""

import os
from docx import Document

def main():
    resume_path = "resume.docx"
    
    if not os.path.exists(resume_path):
        print(f"ERROR: {resume_path} not found!")
        return
    
    doc = Document(resume_path)
    
    # Find the project name paragraph
    project_para = None
    other_project_paras = []
    
    # Find Daily Wellness Scheduler and other project titles
    for para in doc.paragraphs:
        text = para.text.strip()
        if "Daily Wellness Scheduler" in text:
            project_para = para
        elif any(x in text for x in ["G3-GPT", "Security+ Lab Environment", "Doctor Appointment"]):
            other_project_paras.append(para)
    
    if not project_para:
        print("Could not find 'Daily Wellness Scheduler' in resume")
        return
    
    # Get font size from other project titles (like G3-GPT)
    if other_project_paras:
        reference_para = other_project_paras[0]
        if reference_para.runs:
            ref_font_size = reference_para.runs[0].font.size
            ref_font_name = reference_para.runs[0].font.name
            
            # Apply same font size to Daily Wellness Scheduler
            if project_para.runs:
                for run in project_para.runs:
                    if ref_font_size:
                        run.font.size = ref_font_size
                    if ref_font_name:
                        run.font.name = ref_font_name
                    
                    # Also match bold/italic
                    if reference_para.runs[0].font.bold is not None:
                        run.font.bold = reference_para.runs[0].font.bold
                    if reference_para.runs[0].font.italic is not None:
                        run.font.italic = reference_para.runs[0].font.italic
                
                print(f"Fixed font size on 'Daily Wellness Scheduler'")
                print(f"  Matched font size from: {other_project_paras[0].text[:50]}...")
            else:
                # No runs, create one
                run = project_para.add_run(project_para.text)
                if ref_font_size:
                    run.font.size = ref_font_size
                if ref_font_name:
                    run.font.name = ref_font_name
    else:
        print("Could not find other project titles to match font size")
        # Try to set a reasonable default size (11pt is common for resume text)
        from docx.shared import Pt
        if project_para.runs:
            for run in project_para.runs:
                run.font.size = Pt(11)
        print("Set font size to 11pt (default)")
    
    # Save
    doc.save(resume_path)
    print(f"\nResume updated: {resume_path}")
    print("Font size should now match other project titles!")


if __name__ == "__main__":
    main()






