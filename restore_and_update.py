#!/usr/bin/env python3
"""
Restore from backup and update with proper formatting preservation
"""

import sys
import os
import shutil

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("ERROR: python-docx library not found.")
    print("Please install it with: pip install python-docx")
    sys.exit(1)


def main():
    backup_path = "resume_backup.docx"
    resume_path = "resume.docx"
    
    # Check if backup exists
    if not os.path.exists(backup_path):
        print(f"ERROR: Backup file not found: {backup_path}")
        print("Cannot restore. Make sure resume_backup.docx exists.")
        sys.exit(1)
    
    # Check if resume is open (can't overwrite if it is)
    try:
        if os.path.exists(resume_path):
            # Try to remove it
            os.remove(resume_path)
    except PermissionError:
        print("ERROR: Cannot delete resume.docx - it's open in another program.")
        print("Please close Microsoft Word and try again.")
        sys.exit(1)
    
    # Restore from backup
    print(f"Restoring from backup: {backup_path}")
    shutil.copy(backup_path, resume_path)
    print(f"Restored: {resume_path}")
    
    # Now update with formatting preservation
    print(f"\nUpdating resume with proper formatting...")
    doc = Document(resume_path)
    
    changes = []
    bullets_to_add = []
    insert_after_index = None
    insert_after_para = None
    
    # Process all paragraphs - preserve formatting by keeping runs
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # Replace project name - preserve all formatting
        if "Career Development Tracker" in text:
            # Replace text in existing runs to preserve formatting
            if para.runs:
                # Combine all runs into first run, preserving formatting
                first_run = para.runs[0]
                new_text = text.replace("Career Development Tracker", "Daily Wellness Scheduler (Sept 2024 - Present)")
                first_run.text = new_text
                # Remove other runs
                for run in para.runs[1:]:
                    para._element.remove(run._element)
            else:
                para.text = text.replace("Career Development Tracker", "Daily Wellness Scheduler (Sept 2024 - Present)")
            changes.append("Project name updated")
        
        # Handle combined format (both sentences on one line)
        elif "Python desktop application with Tkinter GUI for career planning and task management  Features flexible scheduling" in text:
            # Replace text in first run, preserving formatting
            if para.runs:
                para.runs[0].text = "Python desktop application (2,700+ lines) with tkinter GUI and CLI interface for intelligent supplement scheduling"
                # Remove other runs
                for run in para.runs[1:]:
                    para._element.remove(run._element)
            else:
                para.text = "Python desktop application (2,700+ lines) with tkinter GUI and CLI interface for intelligent supplement scheduling"
            
            insert_after_index = i
            insert_after_para = para
            bullets_to_add = [
                "• Implemented rule-based scheduling algorithm with automatic conflict resolution and time-based rule engine",
                "• Integrated Pushbullet API for real-time phone notifications and missed item detection",
                "• Developed data persistence layer with JSON serialization and export functionality (CSV, iCal)",
                "• Created comprehensive test suite for scheduling logic validation and conflict resolution"
            ]
            changes.append("Combined description updated")
        
        # Handle first description line only
        elif "Python desktop application with Tkinter GUI for career planning and task management" in text and "Features flexible scheduling" not in text:
            # Replace text in runs, preserving formatting
            if para.runs:
                para.runs[0].text = para.runs[0].text.replace(
                    "Python desktop application with Tkinter GUI for career planning and task management",
                    "Python desktop application (2,700+ lines) with tkinter GUI and CLI interface for intelligent supplement scheduling"
                )
                # Remove other runs
                for run in para.runs[1:]:
                    para._element.remove(run._element)
            else:
                para.text = "Python desktop application (2,700+ lines) with tkinter GUI and CLI interface for intelligent supplement scheduling"
            changes.append("Description line 1 updated")
        
        # Handle features line only
        elif "Features flexible scheduling, progress tracking, and GitHub integration" in text:
            # Replace text in first run, preserving formatting
            if para.runs:
                para.runs[0].text = "• Implemented rule-based scheduling algorithm with automatic conflict resolution and time-based rule engine"
                # Remove other runs
                for run in para.runs[1:]:
                    para._element.remove(run._element)
            else:
                para.text = "• Implemented rule-based scheduling algorithm with automatic conflict resolution and time-based rule engine"
            
            insert_after_index = i
            insert_after_para = para
            bullets_to_add = [
                "• Integrated Pushbullet API for real-time phone notifications and missed item detection",
                "• Developed data persistence layer with JSON serialization and export functionality (CSV, iCal)",
                "• Created comprehensive test suite for scheduling logic validation and conflict resolution"
            ]
            changes.append("Features line updated")
    
    # Add bullets after the identified paragraph, preserving formatting
    if insert_after_index is not None and bullets_to_add and insert_after_para:
        # Get formatting from the paragraph we're inserting after
        format_source = insert_after_para.runs[0] if insert_after_para.runs else None
        
        # Insert bullets in reverse order (so they appear in correct order)
        for bullet in reversed(bullets_to_add):
            new_para = doc.paragraphs[insert_after_index].insert_paragraph_before(bullet)
            # Copy formatting from source paragraph
            if format_source:
                # Clear default run and add new one with formatting
                if new_para.runs:
                    new_para.runs[0].text = bullet
                    new_para.runs[0].font.name = format_source.font.name
                    if format_source.font.size:
                        new_para.runs[0].font.size = format_source.font.size
                    if format_source.font.bold is not None:
                        new_para.runs[0].font.bold = format_source.font.bold
    
    # Remove Doctor Appointment App section
    print("\nRemoving Doctor Appointment App section...")
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        para = doc.paragraphs[i]
        if "Doctor Appointment App" in para.text:
            # Remove this paragraph and the next one (description)
            para._element.getparent().remove(para._element)
            if i < len(doc.paragraphs) - 1:
                next_para = doc.paragraphs[i]
                if "Cross-platform Flutter" in next_para.text or "Reduced administrative" in next_para.text:
                    next_para._element.getparent().remove(next_para._element)
            changes.append("Doctor Appointment App section removed")
            break
    
    # Save updated document
    doc.save(resume_path)
    
    print("\n" + "="*60)
    print("RESUME UPDATE COMPLETE")
    print("="*60)
    if changes:
        print("\nChanges made:")
        for change in changes:
            print(f"  - {change}")
    else:
        print("\nNo changes detected")
    
    print(f"\nUpdated file: {resume_path}")
    print("\nFont formatting has been preserved!")
    print("Please open the resume in Word to verify formatting looks correct.")


if __name__ == "__main__":
    main()

